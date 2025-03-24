#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

"""
手册文档解析模块
该模块专门用于处理手册类文档（包括PDF、DOCX等格式）的解析
主要功能包括文档分块、表格识别、图片处理等
"""

import logging
import copy
import re

from api.db import ParserType
from io import BytesIO
from rag.nlp import rag_tokenizer, tokenize, tokenize_table, bullets_category, title_frequency, tokenize_chunks, docx_question_level
from rag.utils import num_tokens_from_string
from deepdoc.parser import PdfParser, PlainParser, DocxParser
from docx import Document
from PIL import Image


class Pdf(PdfParser):
    """PDF手册文档解析器类
    继承自PdfParser，专门用于处理手册类PDF文档
    提供OCR识别、布局分析、表格识别和文本提取等功能
    """
    
    def __init__(self):
        self.model_speciess = ParserType.MANUAL.value
        super().__init__()

    def __call__(self, filename, binary=None, from_page=0,
                 to_page=100000, zoomin=3, callback=None):
        """PDF文档解析的主要方法
        包含以下处理步骤：
        1. OCR识别
        2. 布局分析
        3. 表格识别
        4. 文本合并和清理
        参数:
            filename: 文件名
            binary: 二进制文件内容
            from_page: 起始页码
            to_page: 结束页码
            zoomin: 放大倍数
            callback: 进度回调函数
        返回:
            解析后的文本和表格内容
        """
        from timeit import default_timer as timer
        start = timer()
        callback(msg="OCR started")
        self.__images__(
            filename if not binary else binary,
            zoomin,
            from_page,
            to_page,
            callback
        )
        callback(msg="OCR finished ({:.2f}s)".format(timer() - start))
        logging.debug("OCR: {}".format(timer() - start))

        start = timer()
        self._layouts_rec(zoomin)
        callback(0.65, "Layout analysis ({:.2f}s)".format(timer() - start))
        logging.debug("layouts: {}".format(timer() - start))

        start = timer()
        self._table_transformer_job(zoomin)
        callback(0.67, "Table analysis ({:.2f}s)".format(timer() - start))

        start = timer()
        self._text_merge()
        tbls = self._extract_table_figure(True, zoomin, True, True)
        self._concat_downward()
        self._filter_forpages()
        callback(0.68, "Text merged ({:.2f}s)".format(timer() - start))

        # 清理文本中的多余空白字符
        for b in self.boxes:
            b["text"] = re.sub(r"([\t 　]|\u3000){2,}", " ", b["text"].strip())

        return [(b["text"], b.get("layoutno", ""), self.get_position(b, zoomin))
                for i, b in enumerate(self.boxes)], tbls


class Docx(DocxParser):
    """DOCX手册文档解析器类
    继承自DocxParser，专门用于处理手册类DOCX文档
    提供文档解析、图片处理和表格识别等功能
    """
    
    def __init__(self):
        pass

    def get_picture(self, document, paragraph):
        """从段落中提取图片
        参数:
            document: DOCX文档对象
            paragraph: 段落对象
        返回:
            提取到的图片对象
        """
        img = paragraph._element.xpath('.//pic:pic')
        if not img:
            return None
        img = img[0]
        embed = img.xpath('.//a:blip/@r:embed')[0]
        related_part = document.part.related_parts[embed]
        image = related_part.image
        image = Image.open(BytesIO(image.blob))
        return image

    def concat_img(self, img1, img2):
        """合并两张图片
        将两张图片垂直拼接在一起
        参数:
            img1: 第一张图片
            img2: 第二张图片
        返回:
            合并后的图片对象
        """
        if img1 and not img2:
            return img1
        if not img1 and img2:
            return img2
        if not img1 and not img2:
            return None
        width1, height1 = img1.size
        width2, height2 = img2.size

        new_width = max(width1, width2)
        new_height = height1 + height2
        new_image = Image.new('RGB', (new_width, new_height))

        new_image.paste(img1, (0, 0))
        new_image.paste(img2, (0, height1))

        return new_image

    def __call__(self, filename, binary=None, from_page=0, to_page=100000, callback=None):
        """DOCX文档解析的主要方法
        解析文档并提取问题和答案对，以及表格内容
        参数:
            filename: 文件名
            binary: 二进制文件内容
            from_page: 起始页码
            to_page: 结束页码
            callback: 进度回调函数
        返回:
            解析后的问题答案对和表格内容
        """
        self.doc = Document(
            filename) if not binary else Document(BytesIO(binary))
        pn = 0
        last_answer, last_image = "", None
        question_stack, level_stack = [], []
        ti_list = []
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            question_level, p_text = 0, ''
            if from_page <= pn < to_page and p.text.strip():
                question_level, p_text = docx_question_level(p)
            if not question_level or question_level > 6: # 不是问题
                last_answer = f'{last_answer}\n{p_text}'
                current_image = self.get_picture(self.doc, p)
                last_image = self.concat_img(last_image, current_image)
            else:   # 是问题
                if last_answer or last_image:
                    sum_question = '\n'.join(question_stack)
                    if sum_question:
                        ti_list.append((f'{sum_question}\n{last_answer}', last_image))
                    last_answer, last_image = '', None

                i = question_level
                while question_stack and i <= level_stack[-1]:
                    question_stack.pop()
                    level_stack.pop()
                question_stack.append(p_text)
                level_stack.append(question_level)
            for run in p.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1
                    continue
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        if last_answer:
            sum_question = '\n'.join(question_stack)
            if sum_question:
                ti_list.append((f'{sum_question}\n{last_answer}', last_image))
                
        tbls = []
        for tb in self.doc.tables:
            html= "<table>"
            for r in tb.rows:
                html += "<tr>"
                i = 0
                while i < len(r.cells):
                    span = 1
                    c = r.cells[i]
                    for j in range(i+1, len(r.cells)):
                        if c.text == r.cells[j].text:
                            span += 1
                            i = j
                    i += 1
                    html += f"<td>{c.text}</td>" if span == 1 else f"<td colspan='{span}'>{c.text}</td>"
                html += "</tr>"
            html += "</table>"
            tbls.append(((None, html), ""))
        return ti_list, tbls


def chunk(filename, binary=None, from_page=0, to_page=100000,
          lang="Chinese", callback=None, **kwargs):
    """文档分块处理的主函数
    支持PDF和DOCX格式的手册文档解析
    参数:
        filename: 文件名
        binary: 二进制文件内容
        from_page: 起始页码
        to_page: 结束页码
        lang: 文档语言（默认中文）
        callback: 进度回调函数
        **kwargs: 其他参数
    返回:
        处理后的文档块列表
    """
    pdf_parser = None
    doc = {
        "docnm_kwd": filename
    }
    doc["title_tks"] = rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", doc["docnm_kwd"]))
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    # is it English
    eng = lang.lower() == "english"  # pdf_parser.is_english
    if re.search(r"\.pdf$", filename, re.IGNORECASE):
        pdf_parser = Pdf()
        if kwargs.get("layout_recognize", "DeepDOC") == "Plain Text":
            pdf_parser = PlainParser()
        sections, tbls = pdf_parser(filename if not binary else binary,
                                    from_page=from_page, to_page=to_page, callback=callback)
        if sections and len(sections[0]) < 3:
            sections = [(t, lvl, [[0] * 5]) for t, lvl in sections]
        # 使用最频繁的标题类型设置基准，然后在两个基准之间合并
        if len(sections) > 0 and len(pdf_parser.outlines) / len(sections) > 0.03:
            max_lvl = max([lvl for _, lvl in pdf_parser.outlines])
            most_level = max(0, max_lvl - 1)
            levels = []
            for txt, _, _ in sections:
                for t, lvl in pdf_parser.outlines:
                    tks = set([t[i] + t[i + 1] for i in range(len(t) - 1)])
                    tks_ = set([txt[i] + txt[i + 1]
                                for i in range(min(len(t), len(txt) - 1))])
                    if len(set(tks & tks_)) / max([len(tks), len(tks_), 1]) > 0.8:
                        levels.append(lvl)
                        break
                else:
                    levels.append(max_lvl + 1)

        else:
            bull = bullets_category([txt for txt, _, _ in sections])
            most_level, levels = title_frequency(
                bull, [(txt, lvl) for txt, lvl, _ in sections])

        assert len(sections) == len(levels)
        sec_ids = []
        sid = 0
        for i, lvl in enumerate(levels):
            if lvl <= most_level and i > 0 and lvl != levels[i - 1]:
                sid += 1
            sec_ids.append(sid)

        sections = [(txt, sec_ids[i], poss)
                    for i, (txt, _, poss) in enumerate(sections)]
        for (img, rows), poss in tbls:
            if not rows:
                continue
            sections.append((rows if isinstance(rows, str) else rows[0], -1,
                            [(p[0] + 1 - from_page, p[1], p[2], p[3], p[4]) for p in poss]))

        def tag(pn, left, right, top, bottom):
            """生成位置标签
            参数:
                pn: 页码
                left: 左边距
                right: 右边距
                top: 上边距
                bottom: 下边距
            返回:
                格式化的位置标签字符串
            """
            if pn + left + right + top + bottom == 0:
                return ""
            return "@@{}\t{:.1f}\t{:.1f}\t{:.1f}\t{:.1f}##" \
                .format(pn, left, right, top, bottom)

        chunks = []
        last_sid = -2
        tk_cnt = 0
        for txt, sec_id, poss in sorted(sections, key=lambda x: (
                x[-1][0][0], x[-1][0][3], x[-1][0][1])):
            poss = "\t".join([tag(*pos) for pos in poss])
            if tk_cnt < 32 or (tk_cnt < 1024 and (sec_id == last_sid or sec_id == -1)):
                if chunks:
                    chunks[-1] += "\n" + txt + poss
                    tk_cnt += num_tokens_from_string(txt)
                    continue
            chunks.append(txt + poss)
            tk_cnt = num_tokens_from_string(txt)
            if sec_id > -1:
                last_sid = sec_id

        res = tokenize_table(tbls, doc, eng)
        res.extend(tokenize_chunks(chunks, doc, eng, pdf_parser))
        return res

    elif re.search(r"\.docx?$", filename, re.IGNORECASE):
        docx_parser = Docx()
        ti_list, tbls = docx_parser(filename, binary,
                                    from_page=0, to_page=10000, callback=callback)
        res = tokenize_table(tbls, doc, eng)
        for text, image in ti_list:
            d = copy.deepcopy(doc)
            d['image'] = image
            tokenize(d, text, eng)
            res.append(d)
        return res
    else:
        raise NotImplementedError("file type not supported yet(pdf and docx supported)")
    

if __name__ == "__main__":
    import sys


    def dummy(prog=None, msg=""):
        pass


    chunk(sys.argv[1], callback=dummy)

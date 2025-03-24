#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

"""
法律文档解析模块
该模块提供了对法律文档（包括PDF、DOCX、TXT等格式）的解析功能
主要包含文档分块、文本提取和结构化处理等功能
"""

import logging
from tika import parser
import re
from io import BytesIO
from docx import Document

from api.db import ParserType
from deepdoc.parser.utils import get_text
from rag.nlp import bullets_category, remove_contents_table, hierarchical_merge, \
    make_colon_as_title, tokenize_chunks, docx_question_level
from rag.nlp import rag_tokenizer
from deepdoc.parser import PdfParser, DocxParser, PlainParser, HtmlParser


class Docx(DocxParser):
    """DOCX文档解析器类
    继承自DocxParser，专门用于处理法律文档中的DOCX格式文件
    提供文档分页、文本清理和结构化处理等功能
    """
    
    def __init__(self):
        pass

    def __clean(self, line):
        """清理文本行
        移除全角空格并去除首尾空白字符
        """
        line = re.sub(r"\u3000", " ", line).strip()
        return line

    def old_call(self, filename, binary=None, from_page=0, to_page=100000):
        """旧版本的文档解析方法
        按页解析文档，返回清理后的文本行列表
        """
        self.doc = Document(
            filename) if not binary else Document(BytesIO(binary))
        pn = 0
        lines = []
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            if from_page <= pn < to_page and p.text.strip():
                lines.append(self.__clean(p.text))
            for run in p.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1
                    continue
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        return [line for line in lines if line]

    def __call__(self, filename, binary=None, from_page=0, to_page=100000):
        """文档解析的主要方法
        解析文档并返回结构化的章节内容
        支持分页处理，并保持文档的层级结构
        """
        self.doc = Document(
            filename) if not binary else Document(BytesIO(binary))
        pn = 0
        lines = []
        bull = bullets_category([p.text for p in self.doc.paragraphs])
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            question_level, p_text = docx_question_level(p, bull)
            if not p_text.strip("\n"):
                continue
            lines.append((question_level, p_text))

            for run in p.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1
                    continue
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1

        visit = [False for _ in range(len(lines))]
        sections = []
        for s in range(len(lines)):
            e = s + 1
            while e < len(lines):
                if lines[e][0] <= lines[s][0]:
                    break
                e += 1
            if e - s == 1 and visit[s]:
                continue
            sec = []
            next_level = lines[s][0] + 1
            while not sec and next_level < 22:
                for i in range(s+1, e):
                    if lines[i][0] != next_level:
                        continue
                    sec.append(lines[i][1])
                    visit[i] = True
                next_level += 1
            sec.insert(0, lines[s][1])

            sections.append("\n".join(sec))
        return [s for s in sections if s]

    def __str__(self) -> str:
        """返回对象的字符串表示"""
        return f'''
            question:{self.question},
            answer:{self.answer},
            level:{self.level},
            childs:{self.childs}
        '''


class Pdf(PdfParser):
    """PDF文档解析器类
    继承自PdfParser，专门用于处理法律文档中的PDF格式文件
    提供OCR识别、布局分析和文本提取等功能
    """
    
    def __init__(self):
        self.model_speciess = ParserType.LAWS.value
        super().__init__()

    def __call__(self, filename, binary=None, from_page=0,
                 to_page=100000, zoomin=3, callback=None):
        """PDF文档解析的主要方法
        包含OCR识别、布局分析和文本提取等步骤
        支持进度回调，可以实时反馈处理进度
        """
        from timeit import default_timer as timer
        start = timer()
        callback(msg="OCR started")
        self.__images__(
            filename if not binary else binary,
            zoomin,
            from_page,
            to_page,
            callback
        )
        callback(msg="OCR finished ({:.2f}s)".format(timer() - start))

        start = timer()
        self._layouts_rec(zoomin)
        callback(0.67, "Layout analysis ({:.2f}s)".format(timer() - start))
        logging.debug("layouts:".format(
            ))
        self._naive_vertical_merge()

        callback(0.8, "Text extraction ({:.2f}s)".format(timer() - start))

        return [(b["text"], self._line_tag(b, zoomin))
                for b in self.boxes], None


def chunk(filename, binary=None, from_page=0, to_page=100000,
          lang="Chinese", callback=None, **kwargs):
    """文档分块处理的主函数
    支持多种文档格式（docx、pdf、txt、html、doc）的解析和分块
    参数:
        filename: 文件名
        binary: 二进制文件内容
        from_page: 起始页码
        to_page: 结束页码
        lang: 文档语言（默认中文）
        callback: 进度回调函数
        **kwargs: 其他参数
    返回:
        处理后的文档块列表
    """
    doc = {
        "docnm_kwd": filename,
        "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))
    }
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    pdf_parser = None
    sections = []
    # is it English
    eng = lang.lower() == "english"  # is_english(sections)

    if re.search(r"\.docx$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        chunks = Docx()(filename, binary)
        callback(0.7, "Finish parsing.")
        return tokenize_chunks(chunks, doc, eng, None)

    elif re.search(r"\.pdf$", filename, re.IGNORECASE):
        pdf_parser = Pdf()
        if kwargs.get("layout_recognize", "DeepDOC") == "Plain Text":
            pdf_parser = PlainParser()
        for txt, poss in pdf_parser(filename if not binary else binary,
                                    from_page=from_page, to_page=to_page, callback=callback)[0]:
            sections.append(txt + poss)

    elif re.search(r"\.txt$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        txt = get_text(filename, binary)
        sections = txt.split("\n")
        sections = [s for s in sections if s]
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.(htm|html)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        sections = HtmlParser()(filename, binary)
        sections = [s for s in sections if s]
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.doc$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        binary = BytesIO(binary)
        doc_parsed = parser.from_buffer(binary)
        sections = doc_parsed['content'].split('\n')
        sections = [s for s in sections if s]
        callback(0.8, "Finish parsing.")

    else:
        raise NotImplementedError(
            "file type not supported yet(doc, docx, pdf, txt supported)")

    # Remove 'Contents' part
    remove_contents_table(sections, eng)

    make_colon_as_title(sections)
    bull = bullets_category(sections)
    chunks = hierarchical_merge(bull, sections, 5)
    if not chunks:
        callback(0.99, "No chunk parsed out.")

    return tokenize_chunks(["\n".join(ck)
                           for ck in chunks], doc, eng, pdf_parser)


if __name__ == "__main__":
    import sys

    def dummy(prog=None, msg=""):
        pass
    chunk(sys.argv[1], callback=dummy)
